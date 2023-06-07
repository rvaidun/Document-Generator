#!/usr/bin/env python3
import requests
from bs4 import BeautifulSoup
import bs4
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import nltk
from nltk.tokenize import sent_tokenize
import paraphrase
import wikipedia
import sys
import os
import names
import random
import json
import argparse
import shutil


def get_parser():
    parser = argparse.ArgumentParser(
        formatter_class=argparse.ArgumentDefaultsHelpFormatter
    )
    # add arguments for number of documents to generate and the page to use
    parser.add_argument(
        "-t",
        "--title",
        help="The wikipedia page title to use for generating documents",
        type=str,
        default="Machine Learning",
    )
    parser.add_argument(
        "-n",
        "--number",
        help="The number of documents to generate in a batch",
        default=10,
        type=int,
    )
    parser.add_argument(
        "-b",
        "--batch",
        help="The number of batches to generate",
        default=1,
        type=int,
    )
    # add argument for the class names
    parser.add_argument(
        "-c",
        "--class_names",
        help="The class names to use for generating documents",
        type=list,
        default=[
            "CS 1",
            "CS 2",
            "CS 3",
            "CS 4",
            "CS 5",
            "CS 6",
            "CS 7",
            "CS 8",
            "CS 9",
            "CS 10",
        ],
    )
    parser.add_argument(
        "-s",
        "--sentences",
        help="The number of sentences to use for each chunk",
        default=25,
        type=int,
    )

    return parser.parse_args()


def GenerateBatch(args, batch):
    amt = args.number
    wiki_title = args.title
    num_sentences = args.sentences

    class_names = args.class_names
    generated = 0
    while generated < amt:
        try:
            page = wikipedia.page(wiki_title, auto_suggest=False)
            print("The url for page is " + page.url)
        except Exception as e:
            print("There was an error getting the page. Please try again.")
            print(e)
            sys.exit(1)
        # Split the page into chunks of 5 sentences
        page_content = page.content.split("== See also ==")[0]
        page_content = page_content.strip()
        page_content = sent_tokenize(page_content)

        # Generate chunks
        chunked = []
        count = 0
        curstr = ""
        print("Paraphrasing Chunk " + str(len(chunked)))
        for x in page_content:
            if len(chunked) >= amt:
                break
            if count == num_sentences:
                count = 0
                chunked.append(curstr)
                print("Paraphrasing Chunk " + str(len(chunked)))
                curstr = ""
            curstr += paraphrase.paraphrase(x)
            count += 1
        chunked.append(curstr)
        count = 0
        folder = f"output/{batch} {page.title}"
        if os.path.exists(folder) == False:
            try:
                os.mkdir("./" + folder)
            except OSError:
                print("Creation of the directory failed")

        print("Successfully created the directory")

        # Write chunks to word docs
        for x in chunked:
            print("Writing chunk " + str(count) + " to word doc")
            title = " ".join(nltk.word_tokenize(x)[0:3])
            createDocument(x, title, folder, class_names, doc_no=generated + 1)
            count += 1
            generated += 1
            if generated >= amt:
                break
            else:
                print(
                    "Made "
                    + str(generated)
                    + " docs with this article so far. Need "
                    + str(amt - generated)
                    + " docs more."
                )
    print("Successfully generated " + str(amt) + " documents. Thank you!")


def createDocument(text, title, folder, class_names, doc_no):
    document = Document()
    # Name
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style("NameStyle", WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = "Times New Roman"
    name = document.add_paragraph()
    className = document.add_paragraph()
    name.add_run(names.get_full_name(), style="NameStyle")
    className.add_run(random.choice(class_names), style="NameStyle")
    # Heading
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style("CommentsStyle", WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(16)
    obj_font.name = "Times New Roman"
    paragraph = document.add_paragraph()
    paragraph.add_run(title, style="CommentsStyle").bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Content
    paragraph = document.add_paragraph(text)

    document.save(
        "./"
        + folder
        + "/"
        + f"{doc_no} "
        + "".join(i for i in title.strip() if i not in "\/:*?<>|")
        + ".docx"
    )


if __name__ == "__main__":
    args = get_parser()
    batch = args.batch
    if os.path.exists("output") == True:
        shutil.rmtree("output")
        print("Previous output deleted")
    try:
        os.mkdir("./output")
    except OSError:
        print("Creation of the directory failed")

    print("Successfully created the directory")
    for i in range(batch):
        print("Generating batch " + str(i + 1))
        GenerateBatch(args, i + 1)
